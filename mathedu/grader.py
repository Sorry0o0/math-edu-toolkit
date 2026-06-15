"""
Automated Grading Pipeline for Handwritten Math Answers
Uses multimodal AI (GPT-4o) for vision-based grading.
"""

import base64
import json
from typing import Optional
from dataclasses import dataclass, field


@dataclass
class GradeResult:
    """Result of grading a single problem."""
    number: int
    student_answer: str
    correct_answer: str
    score: float  # 0-10
    feedback: str
    is_correct: bool = False

    def __post_init__(self):
        self.is_correct = self.score >= 8


class Grader:
    """
    Multi-modal grading engine for handwritten math work.
    
    Args:
        model: AI model to use ("gpt-4o", "gpt-4o-mini", "local")
        api_key: OpenAI API key (optional, reads from env if not provided)
    """
    
    def __init__(self, model: str = "gpt-4o", api_key: Optional[str] = None):
        self.model = model
        self.api_key = api_key
        self._client = None
    
    def _get_client(self):
        """Lazy-initialize OpenAI client."""
        if self._client is None:
            try:
                from openai import OpenAI
                self._client = OpenAI(api_key=self.api_key)
            except ImportError:
                raise RuntimeError("openai package not installed. Run: pip install openai")
        return self._client
    
    def grade_from_image(
        self,
        image_path: str,
        answer_key: list,
        config: Optional[dict] = None,
    ) -> list[GradeResult]:
        """
        Grade student answers from a photo of their written work.
        
        Args:
            image_path: Path to the student's answer photo
            answer_key: List of correct answers (strings)
            config: Optional config dict
            
        Returns:
            List of GradeResult objects
        """
        config = config or {}
        
        # Read and encode image
        with open(image_path, "rb") as f:
            image_data = base64.standard_b64encode(f.read()).decode()
        
        # Build prompt with answer key
        key_text = "\n".join(
            f"第{i+1}题：{ans}" for i, ans in enumerate(answer_key)
        )
        
        prompt = f"""你是一位经验丰富的小学数学老师。请批改以下学生手写答案。

## 标准答案
{key_text}

## 批改要求
1. 逐题识别学生的答案
2. 与标准答案比对，每题满分10分
3. 给出具体的反馈（错误原因、鼓励语等）
4. 用中文回复

请以JSON格式返回结果：
{{"results": [{{"number": 1, "student_answer": "...", "score": 10, "feedback": "..."}}, ...]}}
"""
        
        # Call vision API
        client = self._get_client()
        response = client.chat.completions.create(
            model=self.model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{image_data}",
                                "detail": "high",
                            },
                        },
                    ],
                }
            ],
            max_tokens=2000,
            temperature=0.1,
        )
        
        # Parse response
        raw = response.choices[0].message.content
        results = self._parse_grading_response(raw, answer_key)
        return results
    
    def _parse_grading_response(self, raw: str, answer_key: list) -> list[GradeResult]:
        """Parse AI grading response into structured results."""
        try:
            # Try to extract JSON from response
            start = raw.find("{")
            end = raw.rfind("}") + 1
            if start >= 0 and end > start:
                data = json.loads(raw[start:end])
                return [
                    GradeResult(
                        number=r["number"],
                        student_answer=r.get("student_answer", ""),
                        correct_answer=answer_key[r["number"] - 1],
                        score=r.get("score", 0),
                        feedback=r.get("feedback", ""),
                    )
                    for r in data.get("results", [])
                ]
        except (json.JSONDecodeError, KeyError, IndexError):
            pass
        
        # Fallback: return placeholder results
        return [
            GradeResult(
                number=i + 1,
                student_answer="(未识别)",
                correct_answer=ans,
                score=0,
                feedback="无法识别手写内容，请检查图片清晰度",
            )
            for i, ans in enumerate(answer_key)
        ]
    
    def export_feedback(self, results: list[GradeResult], output: str = "feedback.docx"):
        """Export grading results to Word document."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        
        doc = Document()
        doc.add_heading("📝 数学作业批改反馈", 0)
        
        total_score = sum(r.score for r in results)
        max_score = len(results) * 10
        
        # Summary
        p = doc.add_paragraph()
        p.add_run(f"总分：{total_score}/{max_score}  |  ")
        p.add_run(f"正确率：{total_score/max_score*100:.1f}%").bold = True
        
        avg = total_score / len(results) if results else 0
        if avg >= 9:
            comment = "🌟 太棒了！继续保持！"
        elif avg >= 7:
            comment = "👍 做得不错，再接再厉！"
        elif avg >= 5:
            comment = "💪 还需要多加练习哦！"
        else:
            comment = "📚 建议复习相关知识点后再试。"
        doc.add_paragraph(comment)
        
        doc.add_heading("逐题详情", level=1)
        for r in results:
            icon = "✅" if r.is_correct else "❌"
            p = doc.add_paragraph()
            p.add_run(f"{icon} 第{r.number}题：{r.score}/10分\n")
            p.add_run(f"   你的答案：{r.student_answer}\n")
            p.add_run(f"   正确答案：{r.correct_answer}\n")
            run = p.add_run(f"   {r.feedback}")
            if not r.is_correct:
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        
        doc.save(output)
        print(f"✅ Feedback exported to {output}")
