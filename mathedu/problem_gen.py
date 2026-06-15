"""
Problem Generation Engine for Elementary Mathematics
Supports Qingdao Press (青岛版) and other Chinese textbook publishers.
"""

import random
import json
import copy
from typing import Optional
from dataclasses import dataclass, field
from enum import Enum


class Difficulty(Enum):
    EASY = "easy"
    MEDIUM = "medium"
    HARD = "hard"
    MIXED = "mixed"


@dataclass
class MathProblem:
    """Represents a single math problem."""
    number: int
    unit: int
    knowledge_point: str
    difficulty: str
    question_text: str
    answer: str
    solution_steps: list = field(default_factory=list)
    tags: list = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "number": self.number,
            "unit": self.unit,
            "knowledge_point": self.knowledge_point,
            "difficulty": self.difficulty,
            "question": self.question_text,
            "answer": self.answer,
            "steps": self.solution_steps,
            "tags": self.tags,
        }


class ProblemGenerator:
    """
    Generate math problems aligned with curriculum standards.
    
    Args:
        grade: Grade level (1-6)
        publisher: Textbook publisher (qingdao/renjiao/beijing/suke)
        semester: 1 or 2
        unit: Unit number (1-7) or "all"
    """

    # Knowledge points mapped to units (Qingdao Press Grade 3, Sem 2)
    UNIT_KNOWLEDGE_POINTS = {
        1: ["万以上数的认识", "大数的读写", "数的大小比较", "四舍五入"],
        2: ["线的认识", "角的认识与分类", "平行与垂直", "图形的周长"],
        3: ["两位数乘两位数", "乘法估算", "乘法应用题"],
        4: ["混合运算", "两步计算应用题", "运算顺序"],
        5: ["除数是一位数的除法", "除法验算", "有余数的除法", "除法应用题"],
        6: ["分数的初步认识", "分数比较大小", "简单的分数加减法", "分数应用题"],
        7: ["年月日", "24时计时法", "统计与可能性", "排列组合(智慧广场)"],
    }

    def __init__(
        self,
        grade: int = 3,
        publisher: str = "qingdao",
        semester: int = 2,
        unit: Optional[int] = None,
    ):
        self.grade = grade
        self.publisher = publisher
        self.semester = semester
        self.unit = unit
        self._templates = self._load_templates()

    def _load_templates(self) -> dict:
        """Load problem templates from curriculum data."""
        # In production, loads from curriculum_data/
        return _BUILT_IN_TEMPLATES

    def generate(
        self,
        count: int = 20,
        difficulty: Difficulty = Difficulty.MIXED,
        shuffle: bool = True,
    ) -> list[MathProblem]:
        """
        Generate a set of math problems.
        
        Args:
            count: Number of problems to generate
            difficulty: Difficulty level
            shuffle: Whether to randomize order
            
        Returns:
            List of MathProblem objects
        """
        problems = []
        units = [self.unit] if self.unit else list(range(1, 8))
        
        for i in range(count):
            unit = units[i % len(units)]
            kp_list = self.UNIT_KNOWLEDGE_POINTS.get(unit, ["综合"])
            kp = random.choice(kp_list)
            
            diff = (
                difficulty.value
                if difficulty != Difficulty.MIXED
                else random.choice(["easy", "medium", "hard"])
            )
            
            template = self._get_template(unit, kp, diff)
            problem = self._render_template(template, i + 1, unit, kp, diff)
            problems.append(problem)
        
        if shuffle:
            random.shuffle(problems)
            # Re-number after shuffle
            for idx, p in enumerate(problems, 1):
                p.number = idx
        
        return problems

    def _get_template(self, unit: int, kp: str, diff: str) -> dict:
        """Select a problem template based on unit and difficulty."""
        unit_templates = self._templates.get(unit, {})
        kp_templates = unit_templates.get(kp, {})
        templates = kp_templates.get(diff, kp_templates.get("medium", []))
        return random.choice(templates) if templates else self._fallback_template(unit)

    def _render_template(self, template: dict, num: int, unit: int, kp: str, diff: str) -> MathProblem:
        """Render a template into a concrete problem with randomized values."""
        render_fn = template.get("renderer")
        if render_fn:
            result = render_fn()
        else:
            result = self._default_render(template)
        
        return MathProblem(
            number=num,
            unit=unit,
            knowledge_point=kp,
            difficulty=diff,
            question_text=result["question"],
            answer=result["answer"],
            solution_steps=result.get("steps", []),
            tags=result.get("tags", []),
        )

    def _default_render(self, template: dict) -> dict:
        """Default rendering using simple variable substitution."""
        question = template["question_template"]
        answer_template = template["answer_template"]
        
        # Simple numeric substitution
        import re
        vars_needed = re.findall(r'\{(\w+)\}', question)
        values = {}
        for v in vars_needed:
            constraints = template.get("constraints", {}).get(v, {})
            lo, hi = constraints.get("range", [10, 99])
            values[v] = random.randint(lo, hi)
        
        question = question.format(**values)
        answer = eval(answer_template.format(**values)) if '{' in answer_template else answer_template
        
        return {"question": question, "answer": str(answer), "steps": []}

    def _fallback_template(self, unit: int) -> dict:
        """Fallback when no template matches."""
        a, b = random.randint(10, 99), random.randint(2, 9)
        return {
            "question_template": f"计算：{a} × {b} = ？",
            "answer_template": str(a * b),
            "renderer": None,
            "constraints": {},
        }

    def export_pptx(self, problems: list, output: str = "output.pptx"):
        """Export problems to PowerPoint format.
        
        Delegates to the Node.js ppt_generator module via subprocess.
        """
        import subprocess
        import json
        import os
        
        data = [p.to_dict() for p in problems]
        json_path = os.path.join(os.path.dirname(__file__), "_tmp_export.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        script_dir = os.path.join(os.path.dirname(__file__), "..", "ppt_generator")
        subprocess.run(
            ["node", "generate.js", "--input", json_path, "--output", output],
            cwd=script_dir,
            check=True,
        )
        os.remove(json_path)
        print(f"✅ Exported {len(problems)} problems to {output}")

    def export_answers(self, problems: list, output: str = "answers.docx"):
        """Export answer key to Word document format."""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        title = doc.add_heading(f"三年级数学期末复习答案（第1天）", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for p in problems:
            doc.add_paragraph(f"{p.number}. 【{p.knowledge_point}】({p.difficulty})", style="List Number")
            doc.add_paragraph(f"   题目：{p.question_text}")
            pa = doc.add_paragraph(f"   答案：{p.answer}")
            pa.runs[0].font.bold = True
            if p.solution_steps:
                doc.add_paragraph(f"   解答过程：{' → '.join(p.solution_steps)}")
            doc.add_paragraph("")
        
        doc.save(output)
        print(f"✅ Exported answer key to {output}")


# ============================================================
# Built-in Problem Templates (Qingdao Press Grade 3, Sem 2)
# ============================================================

def _gen_multiplication():
    """两位数乘两位数"""
    a = random.randint(11, 49)
    b = random.randint(12, 29)
    steps = [f"{a}×{b}", f"先算{a}×{b%10}={a*(b%10)}", f"再算{a}×{(b//10)*10}={a*(b//10)*10}", f"相加得{a*b}"]
    return {"question": f"用竖式计算：{a} × {b} =", "answer": str(a * b), "steps": steps, "tags": ["竖式计算"]}

def _gen_multiplication_est():
    """乘法估算"""
    a = random.randint(28, 79)
    b = random.randint(28, 59)
    est_a = round(a, -1)
    est_b = round(b, -1)
    return {
        "question": f"估算：{a} × {b} ≈ ？（保留到整十数）",
        "answer": str(est_a * est_b),
        "steps": [f"{a}≈{est_a}", f"{b}≈{est_b}", f"{est_a}×{est_b}={est_a*est_b}"],
        "tags": ["估算"],
    }

def _gen_multiplication_word():
    """乘法应用题"""
    price = random.randint(15, 45)
    qty = random.randint(12, 38)
    total = price * qty
    items = ["书包", "文具盒", "故事书", "篮球", "跳绳"]
    item = random.choice(items)
    return {
        "question": f"学校为每个学生购买一个{item}，每件{price}元，共有{qty}名学生。一共需要多少元？",
        "answer": str(total),
        "steps": [f"单价×数量={price}×{qty}", f"={total}元"],
        "tags": ["应用题", "乘法"],
    }

def _gen_line_angle():
    """线与角"""
    angle_types = [("直角", 90), ("锐角", random.randint(15, 75)), ("钝角", random.randint(95, 165))]
    name, deg = random.choice(angle_types)
    return {
        "question": f"画一个{name}（{deg}°），并标出它的边和顶点。",
        "answer": f"{name}({deg}°)",
        "steps": [f"确定顶点", f"画一条边", f"用量角器量出{deg}°", f"画出另一条边"],
        "tags": ["作图", "角的分类"],
    }

def _gen_perimeter_rect():
    """长方形周长"""
    l = random.randint(5, 25)
    w = random.randint(3, 15)
    return {
        "question": f"一个长方形花坛，长{l}米，宽{w}米。这个花坛的周长是多少米？",
        "answer": str(2 * (l + w)),
        "steps": [f"周长=(长+宽)×2", f"=({l}+{w})×2", f"={2*(l+w)}米"],
        "tags": ["周长", "长方形"],
    }

def _gen_perimeter_square():
    """正方形周长"""
    s = random.randint(4, 18)
    return {
        "question": f"一块正方形手帕，边长{s}厘米。给这块手帕缝一圈花边，需要多长的花边？",
        "answer": str(s * 4),
        "steps": [f"正方形周长=边长×4", f"={s}×4={s*4}厘米"],
        "tags": ["周长", "正方形"],
    }

def _gen_mixed_operation():
    """混合运算"""
    templates = [
        (lambda: (random.randint(10, 50), "+", random.randint(100, 300), "÷", random.randint(2, 9))),
        (lambda: (random.randint(50, 200), "-", random.randint(20, 80), "×", random.randint(3, 7))),
        (lambda: (random.randint(100, 400), "+", random.randint(10, 40), "×", random.randint(5, 9))),
    ]
    fn = random.choice(templates)
    a, op1, b, op2, c = fn()
    expr = f"{a} {op1} {b} {op2} {c}"
    ans = eval(expr)
    return {
        "question": f"脱式计算：{expr}",
        "answer": str(ans),
        "steps": [f"先算{b}{op2}{c}={eval(f'{b}{op2}{c}')}", f"再算{a}{op1}{eval(f'{b}{op2}{c}')}", f"={ans}"],
        "tags": ["脱式计算", "混合运算"],
    }

def _gen_division():
    """除法"""
    divisor = random.choice([2, 3, 4, 5, 6, 7, 8, 9])
    quotient = random.randint(10, 99)
    dividend = divisor * quotient
    remainder = 0
    if random.random() > 0.5:
        remainder = random.randint(1, divisor - 1)
        dividend += remainder
    ans_str = f"{quotient}……{remainder}" if remainder else str(quotient)
    return {
        "question": f"计算：{dividend} ÷ {divisor} = ？",
        "answer": ans_str,
        "steps": [f"{dividend}÷{divisor}", f"商是{quotient}" + (f"，余数是{remainder}" if remainder else "")],
        "tags": ["除法"],
    }

def _gen_fraction_compare():
    """分数比较"""
    denoms = [4, 6, 8, 10, 12]
    d = random.choice(denoms)
    n1, n2 = sorted(random.sample(range(1, d), 2), reverse=True)
    symbol = ">" if n1 > n2 else "<"
    return {
        "question": f"比较大小：{n1}/{d} ○ {n2}/{d}",
        "answer": symbol,
        "steps": [f"分母相同({d})，分子大的分数大", f"{n1}>{n2}，所以{n1}/{d}>{n2}/{d}"],
        "tags": ["分数", "比较大小"],
    }

def _gen_fraction_add():
    """简单分数加法"""
    d = random.choice([4, 6, 8])
    n1 = random.randint(1, d - 2)
    n2 = random.randint(1, d - n1)
    s_num = n1 + n2
    ans = f"{s_num}/{d}" if s_num < d else f"1{'+' + str(s_num-d)+'/'+d if s_num>d else ''}"
    if s_num == d:
        ans = "1"
    elif s_num > d:
        ans = f"{s_num//d}又{s_num%d}/{d}"
    return {
        "question": f"计算：{n1}/{d} + {n2}/{d} = ？",
        "answer": ans,
        "steps": [f"同分母相加，分母不变，分子相加", f"{n1}+{n2}={s_num}", f"={ans}"],
        "tags": ["分数", "加减法"],
    }

def _gen_calendar():
    """年月日"""
    year = random.choice([2024, 2025, 2026, 2027])
    month = random.randint(1, 12)
    days_in_month = [31, 28+(1 if year%4==0 and (year%100!=0 or year%400==0) else 0), 31, 30, 31, 30, 31, 31, 30, 31, 30, 31]
    days = days_in_month[month - 1]
    is_leap = "闰年" if (year % 4 == 0 and (year % 100 != 0 or year % 400 == 0)) else "平年"
    q_type = random.choice(["days_in_month", "leap_year"])
    if q_type == "days_in_month":
        return {
            "question": f"{year}年的{month}月有多少天？这一年是平年还是闰年？",
            "answer": f"{days}天，{is_leap}",
            "steps": [f"{year}%4==0? → {is_leap}", f"{month}月有{days}天"],
            "tags": ["年月日"],
        }
    else:
        return {
            "question": f"{year}年是平年还是闰年？全年有多少天？二月有多少天？",
            "answer": f"{is_leap}，全年{366 if '闰' in is_leap else 365}天，二月{29 if '闰' in is_leap else 28}天",
            "steps": [],
            "tags": ["年月日", "闰年判断"],
        }

def _gen_statistics():
    """统计"""
    categories = ["文学社", "科普组", "历史社", "艺术团"]
    values = [random.randint(5, 18) for _ in range(4)]
    max_idx = values.index(max(values))
    min_idx = values.index(min(values))
    total = sum(values)
    return {
        "question": f"四年级课外小组人数统计：{'、'.join(f'{c}={v}人' for c, v in zip(categories, values))}。\n(1) 哪个小组人数最多？多少人？\n(2) 哪个小组人数最少？多少人？\n(3) 一共多少人？",
        "answer": f"(1){categories[max_idx]}最多，{max(values)}人；(2){categories[min_idx]}最少，{min(values)}人；(3)共{total}人",
        "steps": [f"比较：{values}", f"求和：{'+'.join(map(str, values))}={total}"],
        "tags": ["统计", "数据整理"],
    }

def _gen_24hour_time():
    """24时计时法"""
    h_12 = random.randint(1, 11)
    period = random.choice(["上午", "下午", "晚上"])
    offset = 0 if period == "上午" else 12
    h_24 = h_12 + offset
    if period == "晚上" and h_12 >= 12:
        h_24 = h_12
    direction = random.choice(["to24", "to12"])
    if direction == "to24":
        return {
            "question": f"用24时计时法表示：{period}{h_12}时",
            "answer": f"{h_24}:00 或 {h_24}时",
            "steps": [f"{period}时间+12小时" if offset else "上午时间不变", f"={h_24}时"],
            "tags": ["24时计时法"],
        }
    else:
        h24 = random.randint(13, 23)
        h12 = h24 - 12
        return {
            "question": f"{h24}时是{['下午','晚上'][1 if h24>=18 else 0]}几时？",
            "answer": f"{h12}时（{'晚上' if h24>=18 else '下午'}）",
            "steps": [f"{h24}-12={h12}"],
            "tags": ["24时计时法"],
        }

def _gen_permutation():
    """排列组合（智慧广场）"""
    type_ = random.choice(["permute", "combine"])
    if type_ == "permute":
        items = random.choice([("衣服", 3, "裤子", 2), ("主食", 3, "菜", 3)])
        n1, n2 = items[2], items[3]
        return {
            "question": f"衣柜里有{n1}件不同的{items[1]}和{n2}条不同的{items[3]}，一共可以搭配出多少种不同的穿法？",
            "answer": str(n1 * n2),
            "steps": [f"用乘法原理：{n1}×{n2}={n1*n2}种"],
            "tags": ["排列组合", "智慧广场"],
        }
    else:
        n = random.randint(3, 5)
        r = random.randint(2, min(n, 3))
        from math import perm, comb
        total = perm(n, r)
        names = ["小明", "小红", "小刚", "小丽", "小强"][:n]
        return {
            "question": f"{', '.join(names[:n])}{n}名同学排队照相，{r}人站一排，一共有多少种排法？",
            "answer": str(total),
            "steps": [f"排列：P({n},{r})={n}×{n-1}×...={total}"],
            "tags": ["排列组合", "智慧广场"],
        }

def _gen_word_problem_mixed():
    """综合应用题（BOSS题）"""
    type_ = random.choice(["shopping", "travel", "speed", "distribution"])
    if type_ == "shopping":
        p1, q1 = random.randint(8, 25), random.randint(3, 8)
        p2, q2 = random.randint(15, 45), random.randint(1, 3)
        paid = random.randint(200, 500)
        spent = p1*q1 + p2*q2
        change = paid - spent
        items = ["笔记本", "钢笔"]
        return {
            "question": f"李老师去商店买教学用品，买了{q1}本{items[0]}，每本{p1}元；又买了{q2}支{items[1]}，每支{p2}元。付了{paid}元，应找回多少元？",
            "answer": str(change),
            "steps": [f"{items[0]}:{p1}×{q1}={p1*q1}元", f"{items[1]}:{p2}×{q2}={p2*q2}元", f"共花:{p1*q1+p2*q2}元", f"找回:{paid}-{p1*q1+p2*q2}={change}元"],
            "tags": ["应用题", "综合", "两步计算"],
        }
    elif type_ == "speed":
        speed = random.randint(50, 90)
        time_h = random.randint(2, 6)
        dist = speed * time_h
        return {
            "question": f"一辆汽车以每小时{speed}千米的速度行驶，从A城到B城用了{time_h}小时。A城到B城相距多少千米？如果返回时速度提高{random.randint(10,20)}千米/小时，需要多长时间？",
            "answer": f"{dist}千米；返回需约{dist/(speed+20):.1f}小时",
            "steps": [f"路程=速度×时间={speed}×{time_h}={dist}千米"],
            "tags": ["应用题", "行程问题"],
        }
    elif type_ == "distribution":
        total_items = random.randint(48, 120)
        n_people = random.randint(4, 9)
        per_person = total_items // n_people
        remain = total_items % n_people
        obj = random.choice(["苹果", "铅笔", "糖果", "练习本"])
        return {
            "question": f"把{total_items}个{obj}平均分给{n_people}个小朋友，每人分到几个？还剩几个？",
            "answer": f"{per_person}个，剩{remain}个",
            "steps": [f"{total_items}÷{n_people}={per_person}……{remain}"],
            "tags": ["应用题", "除法", "平均分"],
        }
    else:  # perimeter reverse
        half_p = random.randint(30, 80)
        side_diff = random.randint(3, 15)
        w = (half_p - side_diff) // 2
        l = w + side_diff
        return {
            "question": f"一个长方形的周长是{half_p*2}米，长比宽多{side_diff}米。这个长方形的长和宽各是多少米？面积是多少平方米？",
            "answer": f"长{l}米，宽{w}米，面积{l*w}平方米",
            "steps": [f"长+宽={half_p*2}÷2={half_p}米", f"宽=({half_p}-{side_diff})÷2={w}米", f"长={w}+{side_diff}={l}米", f"面积={l}×{w}={l*w}平方米"],
            "tags": ["应用题", "周长逆推", "综合"],
        }


# Build template registry
_BUILT_IN_TEMPLATES = {
    1: {  # 万以上数的认识
        "大数的读写": {
            "medium": [{"question_template": "{a}", "answer_template": "", "renderer": None}],
        },
    },
    2: {  # 线与角 / 图形的周长
        "角的分类": {
            "easy": [{"renderer": _gen_line_angle}],
            "medium": [{"renderer": _gen_line_angle}],
        },
        "图形的周长": {
            "easy": [
                {"renderer": _gen_perimeter_square},
                {"renderer": _gen_perimeter_rect},
            ],
            "medium": [{"renderer": _gen_perimeter_rect}, {"renderer": _gen_perimeter_square}],
            "hard": [{"renderer": _gen_perimeter_rect}],
        },
    },
    3: {  # 两位数乘两位数
        "乘法口算": {
            "easy": [{"renderer": _gen_multiplication}],
            "medium": [{"renderer": _gen_multiplication}],
        },
        "乘法估算": {
            "easy": [{"renderer": _gen_multiplication_est}],
            "medium": [{"renderer": _gen_multiplication_est}],
        },
        "乘法应用题": {
            "medium": [{"renderer": _gen_multiplication_word}],
            "hard": [{"renderer": _gen_multiplication_word}],
        },
    },
    4: {  # 混合运算
        "脱式计算": {
            "medium": [{"renderer": _gen_mixed_operation}],
            "hard": [{"renderer": _gen_mixed_operation}],
        },
    },
    5: {  # 除数是一位数的除法
        "除法计算": {
            "easy": [{"renderer": _gen_division}],
            "medium": [{"renderer": _gen_division}],
            "hard": [{"renderer": _gen_division}],
        },
    },
    6: {  # 分数的初步认识
        "分数比较": {
            "easy": [{"renderer": _gen_fraction_compare}],
            "medium": [{"renderer": _gen_fraction_compare}],
        },
        "分数加减": {
            "easy": [{"renderer": _gen_fraction_add}],
            "medium": [{"renderer": _gen_fraction_add}],
        },
    },
    7: {  # 年月日 / 统计 / 排列组合
        "年月日": {
            "easy": [{"renderer": _gen_calendar}],
            "medium": [{"renderer": _gen_calendar}],
        },
        "24时计时法": {
            "easy": [{"renderer": _gen_24hour_time}],
            "medium": [{"renderer": _gen_24hour_time}],
        },
        "统计": {
            "easy": [{"renderer": _gen_statistics}],
            "medium": [{"renderer": _gen_statistics}],
        },
        "排列组合": {
            "medium": [{"renderer": _gen_permutation}],
            "hard": [{"renderer": _gen_permutation}],
        },
    },
    0: {  # 综合/BOSS题
        "综合应用": {
            "hard": [{"renderer": _gen_word_problem_mixed}],
            "medium": [{"renderer": _gen_word_problem_mixed}],
        },
    },
}
